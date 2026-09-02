import hashlib
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


MAX_DOCUMENT_BYTES = 10 * 1024 * 1024
MAX_MODEL_CHARACTERS = 100_000
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def extract_document(filename: str, data: bytes) -> tuple[str, str]:
    if not filename:
        raise ValueError("A filename is required")
    extension = Path(filename).suffix.casefold()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError("Supported file types are PDF, DOCX, and TXT")
    if not data:
        raise ValueError("The uploaded document is empty")
    if len(data) > MAX_DOCUMENT_BYTES:
        raise ValueError("The uploaded document exceeds the 10 MB limit")

    if extension == ".pdf":
        reader = PdfReader(BytesIO(data))
        if reader.is_encrypted:
            raise ValueError("Encrypted PDFs are not supported")
        page_parts = []
        for index, page in enumerate(reader.pages, start=1):
            page_text = (page.extract_text() or "").strip()
            if page_text:
                page_parts.append(f"[SOURCE page={index}]\n{page_text}")
        text = "\n\n".join(page_parts)
    elif extension == ".docx":
        document = Document(BytesIO(data))
        parts = [
            f"[SOURCE paragraph={index}]\n{paragraph.text}"
            for index, paragraph in enumerate(document.paragraphs, start=1)
            if paragraph.text.strip()
        ]
        parts.extend(
            f"[SOURCE table={table_index} row={row_index}]\n" + " | ".join(cell.text for cell in row.cells)
            for table_index, table in enumerate(document.tables, start=1)
            for row_index, row in enumerate(table.rows, start=1)
        )
        text = "\n".join(parts)
    else:
        raw_text = data.decode("utf-8-sig", errors="replace")
        lines = raw_text.splitlines() or [raw_text]
        text = "\n\n".join(
            f"[SOURCE lines={start}-{min(start + 39, len(lines))}]\n" +
            "\n".join(lines[start - 1:start + 39])
            for start in range(1, len(lines) + 1, 40)
        )

    text = text.strip()
    if not text:
        raise ValueError(
            "No extractable text was found. Scanned PDFs require OCR, which is not yet supported."
        )
    return text[:MAX_MODEL_CHARACTERS], hashlib.sha256(data).hexdigest()
